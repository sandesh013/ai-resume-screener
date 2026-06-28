"""Extract text from PDF and DOCX files."""
import os

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    if PyPDF2 is None:
        raise ImportError("PyPDF2 is required. Run: pip install PyPDF2")
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text.strip()


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    if docx is None:
        raise ImportError("python-docx is required. Run: pip install python-docx")
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text.strip()


def extract_text(file_path):
    """Extract text based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")